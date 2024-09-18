# import uuid
# from docx import Document
# import os

# def generate_unique_id():
#     """ Generate a unique certificate ID """
#     return str(uuid.uuid4())

# def generate_certificate(name, period):
#     """ Generate and save the certificate based on the template """
#     cert_id = generate_unique_id()
    
#     # Load the DOCX template
#     doc = Document('certificate_template.docx')
    
#     # Replace placeholders with actual values
#     for paragraph in doc.paragraphs:
#         if '[Intern_Name]' in paragraph.text:
#             paragraph.text = paragraph.text.replace('[Intern_Name]', name)
#         if '[Certificate_ID]' in paragraph.text:
#             paragraph.text = paragraph.text.replace('[Certificate_ID]', cert_id)
#         if '[Period]' in paragraph.text:
#             paragraph.text = paragraph.text.replace('[Period]', period)
    
#     # Save the certificate in the certificates folder
#     if not os.path.exists('certificates'):
#         os.makedirs('certificates')
        
#     certificate_path = f"certificates/{cert_id}.docx"
#     doc.save(certificate_path)
    
#     return cert_id, certificate_path




# import uuid
# from docx import Document
# import os

# def generate_unique_id():
#     """Generate a unique certificate ID"""
#     return str(uuid.uuid4())

# def replace_placeholder_in_run(runs, placeholder, replacement):
#     """Replace placeholder text in a run while preserving formatting"""
#     for run in runs:
#         if placeholder in run.text:
#             run.text = run.text.replace(placeholder, replacement)

# def generate_certificate(name, period):
#     """Generate and save the certificate based on the template"""
#     cert_id = generate_unique_id()
    
#     # Load the DOCX template
#     doc = Document('certificate_template.docx')
    
#     # Replace placeholders with actual values while preserving formatting
#     for paragraph in doc.paragraphs:
#         replace_placeholder_in_run(paragraph.runs, '[Intern_Name]', name)
#         replace_placeholder_in_run(paragraph.runs, '[Certificate_ID]', cert_id)
#         replace_placeholder_in_run(paragraph.runs, '[Period]', period)
    
#     # Save the certificate in the certificates folder
#     if not os.path.exists('certificates'):
#         os.makedirs('certificates')
        
#     certificate_path = f"certificates/{cert_id}.docx"
#     doc.save(certificate_path)
    
#     return cert_id, certificate_path



import uuid
from docx import Document
import os

def generate_unique_id():
    """Generate a unique certificate ID"""
    return str(uuid.uuid4())

def replace_placeholder(paragraph, placeholder, replacement):
    """Replace placeholder text in a paragraph, even if split across runs"""
    # Join all the text in the paragraph
    paragraph_text = ''.join([run.text for run in paragraph.runs])
    
    # Replace the placeholder
    if placeholder in paragraph_text:
        # Perform the replacement in the entire paragraph text
        new_text = paragraph_text.replace(placeholder, replacement)
        
        # Clear the current runs and set the new text in a single run
        for run in paragraph.runs:
            run.text = ''  # Clear all existing runs
        paragraph.runs[0].text = new_text  # Set the new text in the first run

def generate_certificate(name, period):
    """Generate and save the certificate based on the template"""
    cert_id = generate_unique_id()
    
    # Load the DOCX template
    doc = Document('certificate_template.docx')
    
    # Replace placeholders in paragraphs while preserving formatting
    for paragraph in doc.paragraphs:
        replace_placeholder(paragraph, '[Intern_Name]', name)
        replace_placeholder(paragraph, '[Certificate_ID]', cert_id)
        replace_placeholder(paragraph, '[Period]', period)
    
    # Save the certificate in the certificates folder
    if not os.path.exists('certificates'):
        os.makedirs('certificates')
        
    certificate_path = f"certificates/{cert_id}.docx"
    doc.save(certificate_path)
    
    return cert_id, certificate_path




